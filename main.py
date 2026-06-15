import csv
import os
import re
import sqlite3
import sys
import tkinter as tk
from dataclasses import dataclass
from datetime import datetime
from tkinter import filedialog, messagebox, ttk
from typing import List, Optional, Tuple


APP_NAME = "Учет должников ТСЖ"
DB_FILE_NAME = "tsj_debtors.db"


def get_app_dir() -> str:
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def resource_path(relative_path: str) -> str:
    if hasattr(sys, "_MEIPASS"):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(get_app_dir(), relative_path)


def get_database_path() -> str:
    return os.path.join(get_app_dir(), DB_FILE_NAME)


def set_window_icon(window) -> None:
    icon_path = resource_path("app.ico")
    if os.path.exists(icon_path):
        try:
            window.iconbitmap(default=icon_path)
        except Exception:
            pass


HELP_EMAIL = "Ivalburashnikov@edu.hse.ru"


def get_help_message() -> str:
    return (
        "Здравствуйте!\n\n"
        "Если возникли вопросы по работе приложения, напишите на почту:\n"
        f"{HELP_EMAIL}"
    )


def prevent_column_resize(event):
    widget = event.widget
    if isinstance(widget, ttk.Treeview):
        if widget.identify_region(event.x, event.y) == "separator":
            return "break"


def safe_filename(text: str) -> str:
    cleaned = re.sub(r'[\\/*?:"<>|]', "_", text)
    return cleaned.strip() or "document"


def format_money(value: float) -> str:
    return f"{value:.2f} руб."


def parse_money(text: str) -> float:
    text = text.strip().replace(",", ".")
    if text == "":
        raise ValueError("Пустое денежное поле.")
    return float(text)


def allow_name_chars(value: str) -> bool:
    return value == "" or re.fullmatch(r"[А-Яа-яЁё\s-]*", value) is not None


def allow_apartment_chars(value: str) -> bool:
    return value == "" or (value.isdigit() and len(value) <= 4)


def allow_personal_account_chars(value: str) -> bool:
    return value == "" or (value.isdigit() and len(value) <= 12)


def allow_phone_chars(value: str) -> bool:
    if value == "":
        return True

    if re.fullmatch(r"[\d\s()+-]*", value) is None:
        return False

    if value.count("+") > 1:
        return False

    if "+" in value and not value.startswith("+"):
        return False

    digits = re.sub(r"\D", "", value)
    return len(digits) <= 11


def allow_money_chars(value: str) -> bool:
    return value == "" or re.fullmatch(r"\d{0,9}([.,]\d{0,2})?", value) is not None


def allow_date_chars(value: str) -> bool:
    return value == "" or (len(value) <= 10 and re.fullmatch(r"[\d.]*", value) is not None)


def allow_period_chars(value: str) -> bool:
    return value == "" or (len(value) <= 15 and re.fullmatch(r"[\d.\-]*", value) is not None)


def validate_full_name(full_name: str) -> Tuple[bool, str]:
    full_name = " ".join(full_name.split())
    pattern = r"^[А-ЯЁ][а-яё-]{1,}\s[А-ЯЁ][а-яё-]{1,}\s[А-ЯЁ][а-яё-]{1,}$"

    if not re.fullmatch(pattern, full_name):
        return False, "ФИО должно состоять из 3 слов на русском языке, например: Иванов Иван Иванович."

    return True, ""


def validate_address(address: str) -> Tuple[bool, str]:
    address = " ".join(address.split())

    if len(address) < 8:
        return False, "Адрес должен быть осмысленным и содержать минимум 8 символов."

    if re.search(r"[А-Яа-яЁё]", address) is None:
        return False, "Адрес должен содержать русские буквы, например: ул. Ленина, д. 10."

    if re.fullmatch(r"[А-Яа-яЁё\s.,\-/№0-9]+", address) is None:
        return False, "Адрес содержит недопустимые символы."

    letters = re.findall(r"[А-Яа-яЁё]", address.lower())
    if len(set(letters)) < 3:
        return False, "Адрес должен быть похож на реальный адрес."

    return True, ""


def validate_apartment_number(apartment_number: str) -> Tuple[bool, str]:
    if not re.fullmatch(r"\d{1,4}", apartment_number.strip()):
        return False, "Номер квартиры должен содержать только цифры."

    return True, ""


def validate_personal_account(personal_account: str) -> Tuple[bool, str]:
    if not re.fullmatch(r"\d{6,12}", personal_account.strip()):
        return False, "Лицевой счёт должен содержать только цифры и иметь длину от 6 до 12 символов."

    return True, ""


def normalize_phone(phone: str) -> Optional[str]:
    digits = re.sub(r"\D", "", phone)

    if len(digits) == 11 and digits.startswith("8"):
        digits = "7" + digits[1:]

    if len(digits) == 11 and digits.startswith("7"):
        return f"+{digits}"

    return None


def validate_phone(phone: str) -> Tuple[bool, str, str]:
    normalized = normalize_phone(phone)

    if normalized is None:
        return False, "", "Телефон должен быть в формате +79991234567, 79991234567, 89991234567 или +7 (999) 123-45-67."

    return True, normalized, ""


def validate_date(date_text: str, allow_empty: bool = False) -> Tuple[bool, str]:
    date_text = date_text.strip()

    if allow_empty and date_text == "":
        return True, ""

    if re.fullmatch(r"\d{2}\.\d{2}\.\d{4}", date_text) is None:
        return False, "Дата должна быть строго в формате ДД.ММ.ГГГГ, например 05.04.2026."

    try:
        datetime.strptime(date_text, "%d.%m.%Y")
        return True, ""
    except ValueError:
        return False, "Такой даты не существует. Проверьте день, месяц и год."


def parse_period(period: str) -> Optional[Tuple[Tuple[int, int], Tuple[int, int]]]:
    match = re.fullmatch(r"(0[1-9]|1[0-2])\.(\d{4})-(0[1-9]|1[0-2])\.(\d{4})", period.strip())

    if not match:
        return None

    start_month = int(match.group(1))
    start_year = int(match.group(2))
    end_month = int(match.group(3))
    end_year = int(match.group(4))

    return (start_year, start_month), (end_year, end_month)


def validate_debt_period(period: str) -> Tuple[bool, str]:
    parsed = parse_period(period)

    if parsed is None:
        return False, "Период задолженности должен быть в формате MM.YYYY-MM.YYYY, например 01.2026-03.2026."

    start, end = parsed

    if start > end:
        return False, "Начало периода задолженности не может быть позже конца."

    return True, ""


def validate_non_negative_number(value: float, field_name: str) -> Tuple[bool, str]:
    if value < 0:
        return False, f"Поле «{field_name}» не может быть отрицательным."

    return True, ""


def create_word_notification(debtor: Tuple, file_path: str) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        raise RuntimeError("Не установлена библиотека python-docx. Выполни команду: python -m pip install python-docx")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("УВЕДОМЛЕНИЕ О ЗАДОЛЖЕННОСТИ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(f"Дата составления: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"ФИО должника: {debtor[1]}")
    doc.add_paragraph(f"Адрес: {debtor[2]}")
    doc.add_paragraph(f"Квартира: {debtor[3]}")
    doc.add_paragraph(f"Лицевой счет: {debtor[4]}")
    doc.add_paragraph(f"Телефон: {debtor[5] or '-'}")
    doc.add_paragraph(f"Период задолженности: {debtor[9] or '-'}")

    doc.add_paragraph(
        f"По состоянию на {datetime.now().strftime('%d.%m.%Y')} за Вами числится задолженность "
        f"по оплате жилищно-коммунальных услуг."
    )

    doc.add_paragraph(f"Основной долг: {format_money(float(debtor[6]))}")
    doc.add_paragraph(f"Пени: {format_money(float(debtor[7]))}")
    doc.add_paragraph(f"Общая сумма задолженности: {format_money(float(debtor[8]))}")

    doc.add_paragraph(
        "Просим Вас в кратчайший срок погасить указанную задолженность. "
        "В случае непогашения долга ТСЖ оставляет за собой право предпринять дальнейшие меры."
    )

    doc.add_paragraph("\nПредседатель ТСЖ ____________________")
    doc.save(file_path)


def create_word_claim(debtor: Tuple, file_path: str) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        raise RuntimeError("Не установлена библиотека python-docx. Выполни команду: python -m pip install python-docx")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ПРЕТЕНЗИЯ О ПОГАШЕНИИ ЗАДОЛЖЕННОСТИ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(f"Дата составления: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"Кому: {debtor[1]}")
    doc.add_paragraph(f"Адрес: {debtor[2]}, кв. {debtor[3]}")
    doc.add_paragraph(f"Лицевой счет: {debtor[4]}")
    doc.add_paragraph(f"Период задолженности: {debtor[9] or '-'}")

    doc.add_paragraph(
        "Настоящим уведомляем Вас о наличии просроченной задолженности перед ТСЖ "
        "по оплате обязательных платежей и коммунальных услуг."
    )

    doc.add_paragraph(f"Основной долг: {format_money(float(debtor[6]))}")
    doc.add_paragraph(f"Пени: {format_money(float(debtor[7]))}")
    doc.add_paragraph(f"Итого к оплате: {format_money(float(debtor[8]))}")

    doc.add_paragraph(
        "Требуем погасить указанную задолженность в добровольном порядке. "
        "При невыполнении данного требования ТСЖ вправе обратиться в суд "
        "для взыскания задолженности, пеней и судебных расходов."
    )

    doc.add_paragraph("\nПредседатель ТСЖ ____________________")
    doc.save(file_path)


@dataclass
class Debtor:
    id: Optional[int]
    full_name: str
    address: str
    apartment_number: str
    personal_account: str
    phone: str
    main_debt: float
    penalty: float
    debt_period: str
    last_payment_date: str
    status: str
    comment: str

    @property
    def total_debt(self) -> float:
        return round(self.main_debt + self.penalty, 2)


class Database:
    def __init__(self, db_name: Optional[str] = None) -> None:
        self.db_name = db_name or get_database_path()
        self._create_tables()
        self._migrate_tables()

    def _connect(self) -> sqlite3.Connection:
        return sqlite3.connect(self.db_name)

    def _create_tables(self) -> None:
        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS debtors (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    full_name TEXT NOT NULL,
                    address TEXT NOT NULL,
                    apartment_number TEXT NOT NULL,
                    personal_account TEXT NOT NULL,
                    phone TEXT,
                    main_debt REAL NOT NULL DEFAULT 0,
                    penalty REAL NOT NULL DEFAULT 0,
                    total_debt REAL NOT NULL DEFAULT 0,
                    debt_period TEXT,
                    last_payment_date TEXT,
                    status TEXT NOT NULL,
                    comment TEXT,
                    initial_main_debt REAL,
                    initial_penalty REAL
                )
                """
            )

            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS payments (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    debtor_id INTEGER NOT NULL,
                    payment_date TEXT NOT NULL,
                    amount REAL NOT NULL,
                    comment TEXT,
                    FOREIGN KEY (debtor_id) REFERENCES debtors(id)
                )
                """
            )

            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS documents (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    debtor_id INTEGER NOT NULL,
                    document_type TEXT NOT NULL,
                    created_at TEXT NOT NULL,
                    file_path TEXT NOT NULL,
                    FOREIGN KEY (debtor_id) REFERENCES debtors(id)
                )
                """
            )

            conn.commit()

    def _migrate_tables(self) -> None:
        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute("PRAGMA table_info(debtors)")
            columns = [row[1] for row in cursor.fetchall()]

            if "initial_main_debt" not in columns:
                cursor.execute("ALTER TABLE debtors ADD COLUMN initial_main_debt REAL")

            if "initial_penalty" not in columns:
                cursor.execute("ALTER TABLE debtors ADD COLUMN initial_penalty REAL")

            cursor.execute(
                """
                UPDATE debtors
                SET initial_main_debt = main_debt
                WHERE initial_main_debt IS NULL
                """
            )

            cursor.execute(
                """
                UPDATE debtors
                SET initial_penalty = penalty
                WHERE initial_penalty IS NULL
                """
            )

            conn.commit()

    def personal_account_exists(self, personal_account: str, exclude_debtor_id: Optional[int] = None) -> bool:
        with self._connect() as conn:
            cursor = conn.cursor()

            if exclude_debtor_id is None:
                cursor.execute(
                    "SELECT COUNT(*) FROM debtors WHERE personal_account = ?",
                    (personal_account,),
                )
            else:
                cursor.execute(
                    "SELECT COUNT(*) FROM debtors WHERE personal_account = ? AND id <> ?",
                    (personal_account, exclude_debtor_id),
                )

            return int(cursor.fetchone()[0]) > 0

    def add_debtor(self, debtor: Debtor) -> int:
        if self.personal_account_exists(debtor.personal_account):
            raise ValueError("Должник с таким лицевым счётом уже есть в базе.")

        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute(
                """
                INSERT INTO debtors (
                    full_name, address, apartment_number, personal_account, phone,
                    main_debt, penalty, total_debt, debt_period, last_payment_date,
                    status, comment, initial_main_debt, initial_penalty
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    debtor.full_name,
                    debtor.address,
                    debtor.apartment_number,
                    debtor.personal_account,
                    debtor.phone,
                    debtor.main_debt,
                    debtor.penalty,
                    debtor.total_debt,
                    debtor.debt_period,
                    debtor.last_payment_date,
                    debtor.status,
                    debtor.comment,
                    debtor.main_debt,
                    debtor.penalty,
                ),
            )

            conn.commit()
            return int(cursor.lastrowid)

    def update_debtor(self, debtor: Debtor) -> None:
        if debtor.id is None:
            raise ValueError("Нельзя обновить должника без ID.")

        if self.personal_account_exists(debtor.personal_account, debtor.id):
            raise ValueError("Должник с таким лицевым счётом уже есть в базе.")

        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute(
                """
                UPDATE debtors
                SET full_name = ?,
                    address = ?,
                    apartment_number = ?,
                    personal_account = ?,
                    phone = ?,
                    debt_period = ?,
                    status = ?,
                    comment = ?,
                    last_payment_date = ?,
                    initial_main_debt = ?,
                    initial_penalty = ?
                WHERE id = ?
                """,
                (
                    debtor.full_name,
                    debtor.address,
                    debtor.apartment_number,
                    debtor.personal_account,
                    debtor.phone,
                    debtor.debt_period,
                    debtor.status,
                    debtor.comment,
                    debtor.last_payment_date,
                    debtor.main_debt,
                    debtor.penalty,
                    debtor.id,
                ),
            )

            conn.commit()

        self.recalculate_debtor_balance(debtor.id)

    def delete_debtor(self, debtor_id: int) -> None:
        with self._connect() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM payments WHERE debtor_id = ?", (debtor_id,))
            cursor.execute("DELETE FROM documents WHERE debtor_id = ?", (debtor_id,))
            cursor.execute("DELETE FROM debtors WHERE id = ?", (debtor_id,))
            conn.commit()

    def get_allowed_payment_amount(self, debtor_id: int, payment_id: Optional[int] = None) -> float:
        debtor = self.get_debtor_by_id(debtor_id)

        if debtor is None:
            raise ValueError("Должник не найден.")

        current_total_debt = float(debtor[8])
        old_amount = 0.0

        if payment_id is not None:
            payment = self.get_payment_by_id(payment_id)

            if payment is not None:
                old_amount = float(payment[3])

        return round(current_total_debt + old_amount, 2)

    def validate_payment_amount(self, debtor_id: int, amount: float, payment_id: Optional[int] = None) -> None:
        if amount <= 0:
            raise ValueError("Сумма оплаты должна быть больше нуля.")

        allowed_amount = self.get_allowed_payment_amount(debtor_id, payment_id)

        if amount > allowed_amount:
            raise ValueError(f"Сумма оплаты не может быть больше допустимого остатка ({allowed_amount:.2f} руб.).")

    def add_payment(self, debtor_id: int, payment_date: str, amount: float, comment: str) -> None:
        self.validate_payment_amount(debtor_id, amount)

        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute(
                """
                INSERT INTO payments (debtor_id, payment_date, amount, comment)
                VALUES (?, ?, ?, ?)
                """,
                (debtor_id, payment_date, amount, comment),
            )

            conn.commit()

        self.recalculate_debtor_balance(debtor_id)

    def update_payment(self, payment_id: int, payment_date: str, amount: float, comment: str) -> None:
        debtor_id = self.get_debtor_id_by_payment(payment_id)

        if debtor_id is None:
            return

        self.validate_payment_amount(debtor_id, amount, payment_id=payment_id)

        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute(
                """
                UPDATE payments
                SET payment_date = ?, amount = ?, comment = ?
                WHERE id = ?
                """,
                (payment_date, amount, comment, payment_id),
            )

            conn.commit()

        self.recalculate_debtor_balance(debtor_id)

    def delete_payment(self, payment_id: int) -> None:
        debtor_id = self.get_debtor_id_by_payment(payment_id)

        if debtor_id is None:
            return

        with self._connect() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM payments WHERE id = ?", (payment_id,))
            conn.commit()

        self.recalculate_debtor_balance(debtor_id)

    def add_document_record(self, debtor_id: int, document_type: str, file_path: str) -> None:
        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute(
                """
                INSERT INTO documents (debtor_id, document_type, created_at, file_path)
                VALUES (?, ?, ?, ?)
                """,
                (
                    debtor_id,
                    document_type,
                    datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                    file_path,
                ),
            )

            conn.commit()

    def get_all_debtors(self) -> List[Tuple]:
        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute(
                """
                SELECT id, full_name, address, apartment_number, personal_account,
                       phone, main_debt, penalty, total_debt, debt_period,
                       last_payment_date, status, comment
                FROM debtors
                ORDER BY id DESC
                """
            )

            return cursor.fetchall()

    def get_filtered_debtors(self, search_text: str = "", status_filter: str = "Все") -> List[Tuple]:
        with self._connect() as conn:
            cursor = conn.cursor()

            query = """
                SELECT id, full_name, address, apartment_number, personal_account,
                       phone, main_debt, penalty, total_debt, debt_period,
                       last_payment_date, status, comment
                FROM debtors
                WHERE 1=1
            """

            params = []

            if search_text:
                query += """
                    AND (
                        full_name LIKE ?
                        OR address LIKE ?
                        OR apartment_number LIKE ?
                        OR personal_account LIKE ?
                        OR phone LIKE ?
                    )
                """

                pattern = f"%{search_text}%"
                params.extend([pattern, pattern, pattern, pattern, pattern])

            if status_filter != "Все":
                query += " AND status = ?"
                params.append(status_filter)

            query += " ORDER BY id DESC"

            cursor.execute(query, params)
            return cursor.fetchall()

    def get_debtor_by_id(self, debtor_id: int) -> Optional[Tuple]:
        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute(
                """
                SELECT id, full_name, address, apartment_number, personal_account,
                       phone, main_debt, penalty, total_debt, debt_period,
                       last_payment_date, status, comment, initial_main_debt, initial_penalty
                FROM debtors
                WHERE id = ?
                """,
                (debtor_id,),
            )

            return cursor.fetchone()

    def get_payments_by_debtor_id(self, debtor_id: int) -> List[Tuple]:
        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute(
                """
                SELECT id, payment_date, amount, comment
                FROM payments
                WHERE debtor_id = ?
                ORDER BY id DESC
                """,
                (debtor_id,),
            )

            return cursor.fetchall()

    def get_payment_by_id(self, payment_id: int) -> Optional[Tuple]:
        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute(
                """
                SELECT id, debtor_id, payment_date, amount, comment
                FROM payments
                WHERE id = ?
                """,
                (payment_id,),
            )

            return cursor.fetchone()

    def get_debtor_id_by_payment(self, payment_id: int) -> Optional[int]:
        row = self.get_payment_by_id(payment_id)

        if row is None:
            return None

        return int(row[1])

    def get_total_paid_by_debtor_id(self, debtor_id: int) -> float:
        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute(
                """
                SELECT COALESCE(SUM(amount), 0)
                FROM payments
                WHERE debtor_id = ?
                """,
                (debtor_id,),
            )

            row = cursor.fetchone()
            return float(row[0] or 0)

    def get_last_payment_date(self, debtor_id: int) -> str:
        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute(
                """
                SELECT payment_date
                FROM payments
                WHERE debtor_id = ?
                ORDER BY id DESC
                LIMIT 1
                """,
                (debtor_id,),
            )

            row = cursor.fetchone()
            return row[0] if row else ""

    def recalculate_debtor_balance(self, debtor_id: int) -> None:
        debtor = self.get_debtor_by_id(debtor_id)

        if debtor is None:
            return

        initial_main = float(debtor[13] or 0)
        initial_penalty = float(debtor[14] or 0)
        current_status = debtor[11]
        manual_last_payment = debtor[10] or ""

        total_paid = self.get_total_paid_by_debtor_id(debtor_id)

        remaining_main = max(0, round(initial_main - total_paid, 2))
        paid_after_main = max(0, round(total_paid - initial_main, 2))
        remaining_penalty = max(0, round(initial_penalty - paid_after_main, 2))
        total_debt = round(remaining_main + remaining_penalty, 2)

        last_payment_date = self.get_last_payment_date(debtor_id) or manual_last_payment

        if total_debt == 0:
            status = "Погашен"
        elif total_paid > 0:
            status = "Частично оплачен"
        else:
            status = current_status if current_status in ("Новый", "Активный") else "Новый"

        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute(
                """
                UPDATE debtors
                SET main_debt = ?,
                    penalty = ?,
                    total_debt = ?,
                    last_payment_date = ?,
                    status = ?
                WHERE id = ?
                """,
                (
                    remaining_main,
                    remaining_penalty,
                    total_debt,
                    last_payment_date,
                    status,
                    debtor_id,
                ),
            )

            conn.commit()

    def get_statistics(self) -> Tuple[int, int, float, float, float, float]:
        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute("SELECT COUNT(*) FROM debtors")
            total_debtors = int(cursor.fetchone()[0])

            cursor.execute("SELECT COUNT(*) FROM debtors WHERE total_debt > 0")
            active_debtors = int(cursor.fetchone()[0])

            cursor.execute(
                """
                SELECT
                    COALESCE(SUM(main_debt), 0),
                    COALESCE(SUM(penalty), 0),
                    COALESCE(SUM(total_debt), 0)
                FROM debtors
                """
            )

            sums = cursor.fetchone()
            sum_main = float(sums[0] or 0)
            sum_penalty = float(sums[1] or 0)
            sum_total = float(sums[2] or 0)

            cursor.execute("SELECT COALESCE(SUM(amount), 0) FROM payments")
            sum_payments = float(cursor.fetchone()[0] or 0)

            return total_debtors, active_debtors, sum_main, sum_penalty, sum_total, sum_payments

    def get_top_debtors(self, limit: int = 10) -> List[Tuple]:
        with self._connect() as conn:
            cursor = conn.cursor()

            cursor.execute(
                """
                SELECT full_name, apartment_number, total_debt, status
                FROM debtors
                ORDER BY total_debt DESC, full_name ASC
                LIMIT ?
                """,
                (limit,),
            )

            return cursor.fetchall()


class DebtorsApp(tk.Tk):
    def __init__(self) -> None:
        super().__init__()

        self.db = Database()

        self.title(APP_NAME)
        set_window_icon(self)

        self.geometry("1280x720")
        self.minsize(1180, 650)

        self.search_var = tk.StringVar()
        self.status_var = tk.StringVar(value="Все")

        self.form_debtor_id: Optional[int] = None
        self.payments_debtor_id: Optional[int] = None
        self.payment_id: Optional[int] = None

        self.form_entries = {}
        self.payment_date_entry = None
        self.payment_amount_entry = None
        self.payment_comment_entry = None

        self.bind_all("<F1>", self.show_context_help)

        self._build_ui()
        self.load_debtors()

    def _build_ui(self) -> None:
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill="both", expand=True)

        self.debtors_tab = ttk.Frame(self.notebook)
        self.form_tab = ttk.Frame(self.notebook)
        self.payments_tab = ttk.Frame(self.notebook)
        self.reports_tab = ttk.Frame(self.notebook)

        self.notebook.add(self.debtors_tab, text="Должники")
        self.notebook.add(self.form_tab, text="Карточка должника")
        self.notebook.add(self.payments_tab, text="Оплаты")
        self.notebook.add(self.reports_tab, text="Отчеты")

        self._build_debtors_tab()
        self._build_form_tab()
        self._build_payments_tab()
        self._build_reports_tab()

    def _build_debtors_tab(self) -> None:
        top_frame = ttk.Frame(self.debtors_tab, padding=10)
        top_frame.pack(fill="x")

        ttk.Label(top_frame, text="Список должников", font=("Arial", 14, "bold")).pack(side="left")

        buttons_frame = ttk.Frame(top_frame)
        buttons_frame.pack(side="right")

        buttons = [
            ("Добавить", self.open_add_debtor),
            ("Редактировать", self.open_edit_debtor),
            ("Оплаты", self.open_payments),
            ("Уведомление", self.create_notification_document),
            ("Претензия", self.create_claim_document),
            ("Отчеты", self.open_reports),
            ("Тестовые данные", self.add_demo_data),
            ("Удалить", self.delete_debtor),
            ("Обновить", self.load_debtors),
            ("Помощь", lambda: self.show_help("index.html")),
        ]

        for text, command in buttons:
            ttk.Button(buttons_frame, text=text, command=command).pack(side="left", padx=4)

        filter_frame = ttk.Frame(self.debtors_tab, padding=(10, 0, 10, 10))
        filter_frame.pack(fill="x")

        ttk.Label(filter_frame, text="Поиск:").pack(side="left", padx=(0, 5))

        search_entry = ttk.Entry(filter_frame, textvariable=self.search_var, width=35)
        search_entry.pack(side="left", padx=(0, 10))
        search_entry.bind("<Return>", lambda event: self.load_debtors())

        ttk.Label(filter_frame, text="Статус:").pack(side="left", padx=(0, 5))

        status_combobox = ttk.Combobox(
            filter_frame,
            textvariable=self.status_var,
            values=["Все", "Новый", "Активный", "Частично оплачен", "Погашен"],
            state="readonly",
            width=18,
        )
        status_combobox.pack(side="left", padx=(0, 10))

        ttk.Button(filter_frame, text="Найти", command=self.load_debtors).pack(side="left", padx=4)
        ttk.Button(filter_frame, text="Сбросить", command=self.reset_filters).pack(side="left", padx=4)

        table_frame = ttk.Frame(self.debtors_tab, padding=(10, 0, 10, 10))
        table_frame.pack(fill="both", expand=True)

        columns = (
            "id",
            "full_name",
            "address",
            "apartment_number",
            "personal_account",
            "phone",
            "main_debt",
            "penalty",
            "total_debt",
            "debt_period",
            "last_payment_date",
            "status",
        )

        self.tree = ttk.Treeview(table_frame, columns=columns, show="headings")
        self.tree.bind("<Double-1>", lambda event: self.open_edit_debtor())
        self.tree.bind("<Button-1>", prevent_column_resize, add="+")

        headings = {
            "id": "ID",
            "full_name": "ФИО",
            "address": "Адрес",
            "apartment_number": "Кв.",
            "personal_account": "Лиц. счет",
            "phone": "Телефон",
            "main_debt": "Осн. долг",
            "penalty": "Пени",
            "total_debt": "Итого",
            "debt_period": "Период",
            "last_payment_date": "Последняя оплата",
            "status": "Статус",
        }

        widths = {
            "id": 50,
            "full_name": 180,
            "address": 190,
            "apartment_number": 60,
            "personal_account": 110,
            "phone": 130,
            "main_debt": 90,
            "penalty": 90,
            "total_debt": 90,
            "debt_period": 140,
            "last_payment_date": 120,
            "status": 130,
        }

        for column in columns:
            self.tree.heading(column, text=headings[column])
            self.tree.column(column, width=widths[column], anchor="center", stretch=False)

        scrollbar_y = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        scrollbar_x = ttk.Scrollbar(table_frame, orient="horizontal", command=self.tree.xview)

        self.tree.configure(yscrollcommand=scrollbar_y.set, xscrollcommand=scrollbar_x.set)

        self.tree.grid(row=0, column=0, sticky="nsew")
        scrollbar_y.grid(row=0, column=1, sticky="ns")
        scrollbar_x.grid(row=1, column=0, sticky="ew")

        table_frame.rowconfigure(0, weight=1)
        table_frame.columnconfigure(0, weight=1)

    def _build_form_tab(self) -> None:
        container = ttk.Frame(self.form_tab, padding=18)
        container.pack(fill="both", expand=True)

        self.form_title_label = ttk.Label(container, text="Карточка должника", font=("Arial", 14, "bold"))
        self.form_title_label.grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 15))

        fields = [
            ("ФИО", "full_name"),
            ("Адрес", "address"),
            ("Номер квартиры", "apartment_number"),
            ("Лицевой счет", "personal_account"),
            ("Телефон", "phone"),
            ("Основной долг", "main_debt"),
            ("Пени", "penalty"),
            ("Период задолженности", "debt_period"),
            ("Дата последней оплаты", "last_payment_date"),
            ("Статус", "status"),
            ("Комментарий", "comment"),
        ]

        for index, (label_text, field_name) in enumerate(fields, start=1):
            ttk.Label(container, text=label_text).grid(row=index, column=0, sticky="w", pady=5)

            if field_name == "comment":
                widget = tk.Text(container, width=60, height=5)
                widget.grid(row=index, column=1, sticky="ew", pady=5)
            elif field_name == "status":
                widget = ttk.Combobox(
                    container,
                    values=["Новый", "Активный", "Частично оплачен", "Погашен"],
                    state="readonly",
                    width=57,
                )
                widget.grid(row=index, column=1, sticky="ew", pady=5)
                widget.set("Новый")
            else:
                widget = ttk.Entry(container, width=60)
                widget.grid(row=index, column=1, sticky="ew", pady=5)
                self._configure_form_validation(field_name, widget)

            self.form_entries[field_name] = widget

        buttons = ttk.Frame(container)
        buttons.grid(row=len(fields) + 1, column=0, columnspan=2, sticky="w", pady=18)

        ttk.Button(buttons, text="Сохранить", command=self.save_debtor_form).pack(side="left", padx=5)
        ttk.Button(buttons, text="Очистить", command=self.clear_debtor_form).pack(side="left", padx=5)
        ttk.Button(buttons, text="Назад к списку", command=lambda: self.notebook.select(self.debtors_tab)).pack(side="left", padx=5)

        container.columnconfigure(1, weight=1)

    def _configure_form_validation(self, field_name: str, widget) -> None:
        validators = {
            "full_name": allow_name_chars,
            "apartment_number": allow_apartment_chars,
            "personal_account": allow_personal_account_chars,
            "phone": allow_phone_chars,
            "main_debt": allow_money_chars,
            "penalty": allow_money_chars,
            "debt_period": allow_period_chars,
            "last_payment_date": allow_date_chars,
        }

        validator = validators.get(field_name)

        if validator is not None:
            widget.configure(validate="key", validatecommand=(self.register(validator), "%P"))

    def _build_payments_tab(self) -> None:
        container = ttk.Frame(self.payments_tab, padding=12)
        container.pack(fill="both", expand=True)

        self.payments_info_label = ttk.Label(container, text="Выберите должника на вкладке «Должники» и нажмите «Оплаты».", font=("Arial", 11, "bold"))
        self.payments_info_label.pack(anchor="w", pady=(0, 10))

        main_area = ttk.Frame(container)
        main_area.pack(fill="both", expand=True)

        table_frame = ttk.LabelFrame(main_area, text="История оплат", padding=10)
        table_frame.pack(side="left", fill="both", expand=True, padx=(0, 10))

        columns = ("id", "payment_date", "amount", "comment")
        self.payments_tree = ttk.Treeview(table_frame, columns=columns, show="headings", height=15)
        self.payments_tree.bind("<Double-1>", lambda event: self.load_selected_payment_to_form())
        self.payments_tree.bind("<Button-1>", prevent_column_resize, add="+")

        self.payments_tree.heading("id", text="ID")
        self.payments_tree.heading("payment_date", text="Дата оплаты")
        self.payments_tree.heading("amount", text="Сумма")
        self.payments_tree.heading("comment", text="Комментарий")

        self.payments_tree.column("id", width=60, anchor="center", stretch=False)
        self.payments_tree.column("payment_date", width=140, anchor="center", stretch=False)
        self.payments_tree.column("amount", width=120, anchor="center", stretch=False)
        self.payments_tree.column("comment", width=330, anchor="w", stretch=False)

        scrollbar_y = ttk.Scrollbar(table_frame, orient="vertical", command=self.payments_tree.yview)
        self.payments_tree.configure(yscrollcommand=scrollbar_y.set)

        self.payments_tree.pack(side="left", fill="both", expand=True)
        scrollbar_y.pack(side="right", fill="y")

        form_frame = ttk.LabelFrame(main_area, text="Оплата", padding=12)
        form_frame.pack(side="right", fill="y")

        ttk.Label(form_frame, text="Дата оплаты").grid(row=0, column=0, sticky="w", pady=5)
        self.payment_date_entry = ttk.Entry(form_frame, width=30)
        self.payment_date_entry.grid(row=0, column=1, sticky="w", pady=5)
        self.payment_date_entry.configure(validate="key", validatecommand=(self.register(allow_date_chars), "%P"))

        ttk.Label(form_frame, text="Сумма оплаты").grid(row=1, column=0, sticky="w", pady=5)
        self.payment_amount_entry = ttk.Entry(form_frame, width=30)
        self.payment_amount_entry.grid(row=1, column=1, sticky="w", pady=5)
        self.payment_amount_entry.configure(validate="key", validatecommand=(self.register(allow_money_chars), "%P"))

        ttk.Label(form_frame, text="Комментарий").grid(row=2, column=0, sticky="w", pady=5)
        self.payment_comment_entry = ttk.Entry(form_frame, width=30)
        self.payment_comment_entry.grid(row=2, column=1, sticky="w", pady=5)

        buttons = ttk.Frame(form_frame)
        buttons.grid(row=3, column=0, columnspan=2, pady=16)

        ttk.Button(buttons, text="Новая", command=self.clear_payment_form).pack(side="left", padx=4)
        ttk.Button(buttons, text="Сохранить", command=self.save_payment_form).pack(side="left", padx=4)
        ttk.Button(buttons, text="Удалить", command=self.delete_payment).pack(side="left", padx=4)

        ttk.Button(form_frame, text="Назад к списку должников", command=lambda: self.notebook.select(self.debtors_tab)).grid(
            row=4, column=0, columnspan=2, sticky="ew", pady=8
        )

    def _build_reports_tab(self) -> None:
        container = ttk.Frame(self.reports_tab, padding=12)
        container.pack(fill="both", expand=True)

        stats_frame = ttk.LabelFrame(container, text="Сводная статистика", padding=10)
        stats_frame.pack(fill="x", pady=(0, 10))

        self.stats_labels = {}

        stat_names = [
            ("total_debtors", "Всего должников"),
            ("active_debtors", "Должников с долгом"),
            ("sum_main", "Общий основной долг"),
            ("sum_penalty", "Общие пени"),
            ("sum_total", "Общая задолженность"),
            ("sum_payments", "Сумма всех оплат"),
        ]

        for index, (key, text) in enumerate(stat_names):
            row = index // 2
            col = (index % 2) * 2

            ttk.Label(stats_frame, text=f"{text}:").grid(row=row, column=col, sticky="w", padx=5, pady=5)

            value_label = ttk.Label(stats_frame, text="0", font=("Arial", 10, "bold"))
            value_label.grid(row=row, column=col + 1, sticky="w", padx=5, pady=5)

            self.stats_labels[key] = value_label

        buttons_frame = ttk.Frame(container)
        buttons_frame.pack(fill="x", pady=(0, 10))

        ttk.Button(buttons_frame, text="Обновить", command=self.refresh_reports).pack(side="left", padx=4)
        ttk.Button(buttons_frame, text="Сохранить CSV", command=self.export_csv).pack(side="left", padx=4)
        ttk.Button(buttons_frame, text="Назад к списку", command=lambda: self.notebook.select(self.debtors_tab)).pack(side="left", padx=4)

        table_frame = ttk.LabelFrame(container, text="Топ должников по сумме долга", padding=10)
        table_frame.pack(fill="both", expand=True)

        columns = ("full_name", "apartment_number", "total_debt", "status")
        self.reports_tree = ttk.Treeview(table_frame, columns=columns, show="headings", height=14)
        self.reports_tree.bind("<Button-1>", prevent_column_resize, add="+")

        self.reports_tree.heading("full_name", text="ФИО")
        self.reports_tree.heading("apartment_number", text="Квартира")
        self.reports_tree.heading("total_debt", text="Сумма долга")
        self.reports_tree.heading("status", text="Статус")

        self.reports_tree.column("full_name", width=320, anchor="w", stretch=False)
        self.reports_tree.column("apartment_number", width=120, anchor="center", stretch=False)
        self.reports_tree.column("total_debt", width=160, anchor="center", stretch=False)
        self.reports_tree.column("status", width=180, anchor="center", stretch=False)

        scrollbar_y = ttk.Scrollbar(table_frame, orient="vertical", command=self.reports_tree.yview)
        self.reports_tree.configure(yscrollcommand=scrollbar_y.set)

        self.reports_tree.pack(side="left", fill="both", expand=True)
        scrollbar_y.pack(side="right", fill="y")

    def show_help(self, topic: str = "index.html") -> None:
        messagebox.showinfo("Помощь", get_help_message())

    def show_context_help(self, event=None) -> str:
        tab_index = self.notebook.index(self.notebook.select())

        if tab_index == 0:
            topic = "main_window.html"
        elif tab_index == 1:
            topic = "debtor_form.html"
        elif tab_index == 2:
            topic = "payments.html"
        else:
            topic = "reports.html"

        self.show_help(topic)
        return "break"

    def reset_filters(self) -> None:
        self.search_var.set("")
        self.status_var.set("Все")
        self.load_debtors()

    def load_debtors(self) -> None:
        for item in self.tree.get_children():
            self.tree.delete(item)

        rows = self.db.get_filtered_debtors(
            search_text=self.search_var.get().strip(),
            status_filter=self.status_var.get().strip(),
        )

        for row in rows:
            self.tree.insert("", tk.END, values=row[:-1])

    def _get_selected_debtor_id(self) -> Optional[int]:
        selected = self.tree.selection()

        if not selected:
            return None

        values = self.tree.item(selected[0], "values")
        return int(values[0])

    def _get_selected_debtor(self) -> Optional[Tuple]:
        debtor_id = self._get_selected_debtor_id()

        if debtor_id is None:
            return None

        return self.db.get_debtor_by_id(debtor_id)

    def _get_text_value(self, field_name: str) -> str:
        widget = self.form_entries[field_name]

        if isinstance(widget, tk.Text):
            return widget.get("1.0", tk.END).strip()

        return widget.get().strip()

    def _set_text_value(self, field_name: str, value: str) -> None:
        widget = self.form_entries[field_name]

        if isinstance(widget, tk.Text):
            widget.delete("1.0", tk.END)
            widget.insert("1.0", value or "")
        elif isinstance(widget, ttk.Combobox):
            widget.set(value or "Новый")
        else:
            widget.delete(0, tk.END)
            widget.insert(0, value or "")

    def clear_debtor_form(self) -> None:
        self.form_debtor_id = None
        self.form_title_label.config(text="Добавление должника")

        for field_name, widget in self.form_entries.items():
            if isinstance(widget, tk.Text):
                widget.delete("1.0", tk.END)
            elif isinstance(widget, ttk.Combobox):
                widget.set("Новый")
            else:
                widget.delete(0, tk.END)

    def open_add_debtor(self) -> None:
        self.clear_debtor_form()
        self.notebook.select(self.form_tab)

    def open_edit_debtor(self) -> None:
        debtor_id = self._get_selected_debtor_id()

        if debtor_id is None:
            messagebox.showwarning("Выбор", "Сначала выберите запись для редактирования.")
            return

        row = self.db.get_debtor_by_id(debtor_id)

        if row is None:
            messagebox.showerror("Ошибка", "Должник не найден.")
            return

        self.form_debtor_id = debtor_id
        self.form_title_label.config(text="Редактирование должника")

        values = {
            "full_name": row[1],
            "address": row[2],
            "apartment_number": row[3],
            "personal_account": row[4],
            "phone": row[5],
            "main_debt": str(row[13] if row[13] is not None else row[6]),
            "penalty": str(row[14] if row[14] is not None else row[7]),
            "debt_period": row[9],
            "last_payment_date": row[10],
            "status": row[11],
            "comment": row[12],
        }

        for field_name, value in values.items():
            self._set_text_value(field_name, value)

        self.notebook.select(self.form_tab)

    def save_debtor_form(self) -> None:
        try:
            full_name = " ".join(self._get_text_value("full_name").split())
            address = self._get_text_value("address")
            apartment_number = self._get_text_value("apartment_number")
            personal_account = self._get_text_value("personal_account")
            phone = self._get_text_value("phone")
            main_debt_text = self._get_text_value("main_debt")
            penalty_text = self._get_text_value("penalty")
            debt_period = self._get_text_value("debt_period")
            last_payment_date = self._get_text_value("last_payment_date")
            status = self._get_text_value("status") or "Новый"
            comment = self._get_text_value("comment")

            if main_debt_text == "" or penalty_text == "":
                messagebox.showwarning("Проверка", "Поля «Основной долг» и «Пени» должны быть заполнены.")
                return

            main_debt = parse_money(main_debt_text)
            penalty = parse_money(penalty_text)

        except ValueError:
            messagebox.showerror("Ошибка", "Поля долга и пени должны быть числами.")
            return

        checks = [
            validate_full_name(full_name),
            validate_address(address),
            validate_apartment_number(apartment_number),
            validate_personal_account(personal_account),
            validate_debt_period(debt_period),
            validate_date(last_payment_date, allow_empty=True),
            validate_non_negative_number(main_debt, "Основной долг"),
            validate_non_negative_number(penalty, "Пени"),
        ]

        for is_valid, error_message in checks:
            if not is_valid:
                messagebox.showwarning("Проверка", error_message)
                return

        phone_valid, normalized_phone, phone_error = validate_phone(phone)

        if not phone_valid:
            messagebox.showwarning("Проверка", phone_error)
            return

        debtor = Debtor(
            id=self.form_debtor_id,
            full_name=full_name,
            address=address.strip(),
            apartment_number=apartment_number.strip(),
            personal_account=personal_account.strip(),
            phone=normalized_phone,
            main_debt=main_debt,
            penalty=penalty,
            debt_period=debt_period.strip(),
            last_payment_date=last_payment_date.strip(),
            status=status,
            comment=comment,
        )

        try:
            if self.form_debtor_id is None:
                self.db.add_debtor(debtor)
            else:
                self.db.update_debtor(debtor)
        except ValueError as error:
            messagebox.showwarning("Проверка", str(error))
            return

        self.load_debtors()
        self.refresh_reports()
        messagebox.showinfo("Готово", "Данные должника сохранены.")
        self.notebook.select(self.debtors_tab)

    def open_payments(self) -> None:
        debtor_id = self._get_selected_debtor_id()

        if debtor_id is None:
            messagebox.showwarning("Выбор", "Сначала выберите должника.")
            return

        self.payments_debtor_id = debtor_id
        self.clear_payment_form()
        self.refresh_payments()
        self.notebook.select(self.payments_tab)

    def refresh_payments(self) -> None:
        for item in self.payments_tree.get_children():
            self.payments_tree.delete(item)

        if self.payments_debtor_id is None:
            self.payments_info_label.config(text="Выберите должника на вкладке «Должники» и нажмите «Оплаты».")
            return

        debtor = self.db.get_debtor_by_id(self.payments_debtor_id)

        if debtor is None:
            self.payments_info_label.config(text="Должник не найден.")
            return

        self.payments_info_label.config(
            text=(
                f"{debtor[1]} | Осн. долг: {debtor[6]:.2f} | "
                f"Пени: {debtor[7]:.2f} | Итого: {debtor[8]:.2f}"
            )
        )

        payments = self.db.get_payments_by_debtor_id(self.payments_debtor_id)

        for payment in payments:
            self.payments_tree.insert("", tk.END, values=payment)

    def clear_payment_form(self) -> None:
        self.payment_id = None

        if self.payment_date_entry is not None:
            self.payment_date_entry.delete(0, tk.END)

        if self.payment_amount_entry is not None:
            self.payment_amount_entry.delete(0, tk.END)

        if self.payment_comment_entry is not None:
            self.payment_comment_entry.delete(0, tk.END)

    def _get_selected_payment_id(self) -> Optional[int]:
        selected = self.payments_tree.selection()

        if not selected:
            return None

        values = self.payments_tree.item(selected[0], "values")
        return int(values[0])

    def load_selected_payment_to_form(self) -> None:
        payment_id = self._get_selected_payment_id()

        if payment_id is None:
            messagebox.showwarning("Выбор", "Сначала выберите оплату.")
            return

        payment = self.db.get_payment_by_id(payment_id)

        if payment is None:
            messagebox.showerror("Ошибка", "Оплата не найдена.")
            return

        self.payment_id = payment_id

        self.payment_date_entry.delete(0, tk.END)
        self.payment_date_entry.insert(0, payment[2])

        self.payment_amount_entry.delete(0, tk.END)
        self.payment_amount_entry.insert(0, str(payment[3]))

        self.payment_comment_entry.delete(0, tk.END)
        self.payment_comment_entry.insert(0, payment[4] or "")

    def save_payment_form(self) -> None:
        if self.payments_debtor_id is None:
            messagebox.showwarning("Выбор", "Сначала выберите должника на вкладке «Должники».")
            return

        payment_date = self.payment_date_entry.get().strip()
        amount_text = self.payment_amount_entry.get().strip()
        comment = self.payment_comment_entry.get().strip()

        if amount_text == "":
            messagebox.showwarning("Проверка", "Сумма оплаты должна быть заполнена.")
            return

        try:
            amount = parse_money(amount_text)
        except ValueError:
            messagebox.showerror("Ошибка", "Сумма оплаты должна быть числом.")
            return

        is_valid_date, date_error = validate_date(payment_date, allow_empty=False)

        if not is_valid_date:
            messagebox.showwarning("Проверка", date_error)
            return

        is_valid_amount, amount_error = validate_non_negative_number(amount, "Сумма оплаты")

        if not is_valid_amount:
            messagebox.showwarning("Проверка", amount_error)
            return

        try:
            self.db.validate_payment_amount(self.payments_debtor_id, amount, payment_id=self.payment_id)

            if self.payment_id is None:
                self.db.add_payment(self.payments_debtor_id, payment_date, amount, comment)
            else:
                self.db.update_payment(self.payment_id, payment_date, amount, comment)

        except ValueError as error:
            messagebox.showwarning("Проверка", str(error))
            return

        self.clear_payment_form()
        self.refresh_payments()
        self.load_debtors()
        self.refresh_reports()
        messagebox.showinfo("Готово", "Оплата сохранена.")

    def delete_payment(self) -> None:
        payment_id = self._get_selected_payment_id()

        if payment_id is None:
            messagebox.showwarning("Выбор", "Сначала выберите оплату.")
            return

        confirmed = messagebox.askyesno("Подтверждение", "Удалить выбранную оплату?")

        if not confirmed:
            return

        self.db.delete_payment(payment_id)
        self.clear_payment_form()
        self.refresh_payments()
        self.load_debtors()
        self.refresh_reports()

    def open_reports(self) -> None:
        self.refresh_reports()
        self.notebook.select(self.reports_tab)

    def refresh_reports(self) -> None:
        total_debtors, active_debtors, sum_main, sum_penalty, sum_total, sum_payments = self.db.get_statistics()

        self.stats_labels["total_debtors"].config(text=str(total_debtors))
        self.stats_labels["active_debtors"].config(text=str(active_debtors))
        self.stats_labels["sum_main"].config(text=f"{sum_main:.2f}")
        self.stats_labels["sum_penalty"].config(text=f"{sum_penalty:.2f}")
        self.stats_labels["sum_total"].config(text=f"{sum_total:.2f}")
        self.stats_labels["sum_payments"].config(text=f"{sum_payments:.2f}")

        for item in self.reports_tree.get_children():
            self.reports_tree.delete(item)

        for row in self.db.get_top_debtors():
            self.reports_tree.insert("", tk.END, values=row)

    def export_csv(self) -> None:
        file_path = filedialog.asksaveasfilename(
            title="Сохранить отчет",
            defaultextension=".csv",
            filetypes=[("CSV файлы", "*.csv")],
        )

        if not file_path:
            return

        rows = self.db.get_all_debtors()

        with open(file_path, "w", newline="", encoding="utf-8-sig") as file:
            writer = csv.writer(file, delimiter=";")

            writer.writerow(
                [
                    "ID",
                    "ФИО",
                    "Адрес",
                    "Квартира",
                    "Лицевой счет",
                    "Телефон",
                    "Основной долг",
                    "Пени",
                    "Итого",
                    "Период задолженности",
                    "Дата последней оплаты",
                    "Статус",
                    "Комментарий",
                ]
            )

            writer.writerows(rows)

        messagebox.showinfo("Готово", "Отчет сохранен в CSV.")

    def create_notification_document(self) -> None:
        debtor = self._get_selected_debtor()

        if debtor is None:
            messagebox.showwarning("Выбор", "Сначала выберите должника.")
            return

        default_name = f"Уведомление_{safe_filename(debtor[1])}_{datetime.now().strftime('%Y-%m-%d')}.docx"

        file_path = filedialog.asksaveasfilename(
            title="Сохранить уведомление",
            defaultextension=".docx",
            initialfile=default_name,
            filetypes=[("Word document", "*.docx")],
        )

        if not file_path:
            return

        try:
            create_word_notification(debtor, file_path)
            self.db.add_document_record(int(debtor[0]), "Уведомление", file_path)
            messagebox.showinfo("Готово", "Уведомление успешно создано.")
        except Exception as error:
            messagebox.showerror("Ошибка", str(error))

    def create_claim_document(self) -> None:
        debtor = self._get_selected_debtor()

        if debtor is None:
            messagebox.showwarning("Выбор", "Сначала выберите должника.")
            return

        default_name = f"Претензия_{safe_filename(debtor[1])}_{datetime.now().strftime('%Y-%m-%d')}.docx"

        file_path = filedialog.asksaveasfilename(
            title="Сохранить претензию",
            defaultextension=".docx",
            initialfile=default_name,
            filetypes=[("Word document", "*.docx")],
        )

        if not file_path:
            return

        try:
            create_word_claim(debtor, file_path)
            self.db.add_document_record(int(debtor[0]), "Претензия", file_path)
            messagebox.showinfo("Готово", "Претензия успешно создана.")
        except Exception as error:
            messagebox.showerror("Ошибка", str(error))

    def add_demo_data(self) -> None:
        existing_rows = self.db.get_all_debtors()

        if existing_rows:
            confirmed = messagebox.askyesno(
                "Тестовые данные",
                "В базе уже есть записи. Добавить тестовые данные поверх существующих?"
            )

            if not confirmed:
                return

        demo_debtors = [
            Debtor(None, "Иванов Иван Иванович", "ул. Ленина, д. 10", "12", "123456", "+79001112233", 12000, 350, "01.2026-03.2026", "", "Новый", "Первичное уведомление"),
            Debtor(None, "Петров Петр Сергеевич", "ул. Мира, д. 5", "24", "234567", "+79002223344", 8500, 200, "02.2026-03.2026", "", "Активный", ""),
            Debtor(None, "Сидорова Анна Викторовна", "ул. Гагарина, д. 8", "7", "345678", "+79003334455", 6700, 150, "12.2025-03.2026", "", "Новый", ""),
            Debtor(None, "Кузнецов Алексей Игоревич", "ул. Центральная, д. 3", "41", "456789", "+79004445566", 15000, 500, "11.2025-03.2026", "", "Активный", "Длительная задолженность"),
            Debtor(None, "Морозова Елена Павловна", "ул. Школьная, д. 15", "9", "567890", "+79005556677", 4300, 120, "03.2026-03.2026", "", "Новый", ""),
            Debtor(None, "Федоров Николай Андреевич", "ул. Садовая, д. 21", "18", "678901", "+79006667788", 9800, 260, "01.2026-03.2026", "", "Активный", ""),
        ]

        created_ids = []

        for debtor in demo_debtors:
            try:
                debtor_id = self.db.add_debtor(debtor)
                created_ids.append(debtor_id)
            except ValueError:
                pass

        if not created_ids:
            messagebox.showinfo("Тестовые данные", "Тестовые данные уже есть в базе.")
            return

        demo_payments = []

        if len(created_ids) > 0:
            demo_payments.append((created_ids[0], "05.04.2026", 3000, "Частичная оплата"))

        if len(created_ids) > 1:
            demo_payments.append((created_ids[1], "02.04.2026", 8700, "Полное погашение"))

        if len(created_ids) > 3:
            demo_payments.append((created_ids[3], "08.04.2026", 5000, "Частичная оплата"))

        if len(created_ids) > 5:
            demo_payments.append((created_ids[5], "09.04.2026", 2000, "Частичная оплата"))

        for debtor_id, payment_date, amount, comment in demo_payments:
            try:
                self.db.add_payment(debtor_id, payment_date, amount, comment)
            except ValueError:
                pass

        self.load_debtors()
        self.refresh_reports()
        messagebox.showinfo("Готово", "Тестовые данные успешно добавлены.")

    def delete_debtor(self) -> None:
        debtor_id = self._get_selected_debtor_id()

        if debtor_id is None:
            messagebox.showwarning("Выбор", "Сначала выберите запись для удаления.")
            return

        confirmed = messagebox.askyesno("Подтверждение", "Удалить выбранного должника?")

        if not confirmed:
            return

        self.db.delete_debtor(debtor_id)
        self.load_debtors()
        self.refresh_reports()


if __name__ == "__main__":
    app = DebtorsApp()
    app.mainloop()